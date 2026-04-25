"""
Outil : exports vers fichiers téléchargeables (PDF / DOCX / XLSX).

PDF  : reportlab     (factures officielles, attestations)
DOCX : python-docx   (lettres modifiables)
XLSX : openpyxl      (exports de données)

Tous les fichiers sont créés dans `exports/` et le chemin est renvoyé
dans le résultat avec le marqueur '📎 FICHIER :' que Streamlit détecte
pour afficher un bouton de téléchargement.
"""
from __future__ import annotations

import csv
from datetime import datetime, date
from pathlib import Path

from .data_processor import _charger_employes
from .facturation import _trouver_facture, _trouver_client, ENTREPRISE, _fmt as _fmt_chf

EXPORTS_DIR = Path(__file__).parent.parent / "exports"
EXPORTS_DIR.mkdir(exist_ok=True)


def _marqueur(path: Path) -> str:
    """Marqueur lu par l'UI Streamlit pour afficher un bouton de téléchargement."""
    return f"\n\n📎 FICHIER : {path.absolute()}"


# =====================================================================
# PDF — factures officielles (avec reportlab)
# =====================================================================
def exporter_facture_pdf(numero: str) -> str:
    """Génère un PDF officiel d'une facture existante."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        )
    except ImportError:
        return ("❌ Le module 'reportlab' n'est pas installé.\n"
                "Lance : pip install reportlab")

    f = _trouver_facture(numero)
    if not f:
        return f"Facture introuvable : {numero}"
    client = _trouver_client(f["id_client"])
    if not client:
        return f"Client introuvable pour {numero}"

    pdf_path = EXPORTS_DIR / f"{numero}.pdf"
    doc = SimpleDocTemplate(
        str(pdf_path), pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    style_t = ParagraphStyle("titre", parent=styles["Title"], fontSize=20, spaceAfter=14)
    style_n = ParagraphStyle("normal", parent=styles["Normal"], fontSize=10, leading=14)

    story = []

    # En-tête entreprise
    story.append(Paragraph(f"<b>{ENTREPRISE['nom']}</b>", style_n))
    story.append(Paragraph(
        f"{ENTREPRISE['adresse']}<br/>{ENTREPRISE['npa']} {ENTREPRISE['ville']}<br/>"
        f"N° TVA : {ENTREPRISE['no_tva']}",
        style_n,
    ))
    story.append(Spacer(1, 0.5 * cm))

    # Titre
    story.append(Paragraph(f"FACTURE N° {numero}", style_t))

    # Bloc destinataire
    story.append(Paragraph("<b>Adressée à :</b>", style_n))
    story.append(Paragraph(
        f"{client['nom']}<br/>{client['contact']}<br/>"
        f"{client['adresse']}<br/>{client['npa']} {client['ville']}<br/>"
        f"N° TVA : {client.get('no_tva', '-')}",
        style_n,
    ))
    story.append(Spacer(1, 0.5 * cm))

    # Dates
    d_em = datetime.strptime(f["date_emission"], "%Y-%m-%d").strftime("%d/%m/%Y")
    d_ec = datetime.strptime(f["date_echeance"], "%Y-%m-%d").strftime("%d/%m/%Y")
    story.append(Paragraph(
        f"<b>Date d'émission :</b> {d_em}<br/><b>Date d'échéance :</b> {d_ec}",
        style_n,
    ))
    story.append(Spacer(1, 0.5 * cm))

    # Tableau prestations
    data = [
        ["Description", "Montant HT (CHF)"],
        [f["description"], _fmt_chf(float(f["montant_ht"]))],
        ["", ""],
        [f"TVA ({f['taux_tva']} %)", _fmt_chf(float(f["montant_tva"]))],
        ["TOTAL TTC", _fmt_chf(float(f["montant_ttc"]))],
    ]
    table = Table(data, colWidths=[12 * cm, 4 * cm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4F46E5")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#EEF2FF")),
        ("ALIGN", (1, 0), (1, -1), "RIGHT"),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
        ("PADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(table)
    story.append(Spacer(1, 0.8 * cm))

    # Coordonnées de paiement
    story.append(Paragraph("<b>Coordonnées de paiement</b>", style_n))
    story.append(Paragraph(
        f"IBAN : {ENTREPRISE['iban']}<br/>"
        f"BIC : {ENTREPRISE['bic']}<br/>"
        f"Référence à indiquer : <b>{numero}</b>",
        style_n,
    ))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph(
        f"<i>Pour toute question : {ENTREPRISE['email']} / {ENTREPRISE['telephone']}</i>",
        style_n,
    ))

    doc.build(story)
    return f"✅ PDF généré : {pdf_path.name}{_marqueur(pdf_path)}"


# =====================================================================
# PDF — attestation officielle
# =====================================================================
def exporter_attestation_pdf(matricule: str, motif: str = "") -> str:
    """Génère une attestation de travail au format PDF officiel."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except ImportError:
        return "❌ Lance : pip install reportlab"

    emp = next((e for e in _charger_employes() if e["matricule"].upper() == matricule.upper()), None)
    if not emp:
        return f"Matricule introuvable : {matricule}"

    pdf_path = EXPORTS_DIR / f"attestation_{matricule}_{date.today().strftime('%Y%m%d')}.pdf"
    doc = SimpleDocTemplate(str(pdf_path), pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2.5*cm, bottomMargin=2.5*cm)
    styles = getSampleStyleSheet()
    sn = ParagraphStyle("n", parent=styles["Normal"], fontSize=11, leading=16)
    st = ParagraphStyle("t", parent=styles["Title"], fontSize=18, alignment=1, spaceAfter=20)

    story = [
        Paragraph(f"<b>{ENTREPRISE['nom']}</b>", sn),
        Paragraph(f"{ENTREPRISE['adresse']}<br/>{ENTREPRISE['npa']} {ENTREPRISE['ville']}", sn),
        Spacer(1, 0.5 * cm),
        Paragraph(f"<i>Fait à {ENTREPRISE['ville']}, le {date.today().strftime('%d/%m/%Y')}</i>", sn),
        Spacer(1, 1.5 * cm),
        Paragraph("ATTESTATION DE TRAVAIL", st),
        Spacer(1, 0.5 * cm),
        Paragraph(
            f"Je soussignée, en qualité de Directrice des Ressources Humaines de la société "
            f"<b>{ENTREPRISE['nom']}</b>, atteste par la présente que :",
            sn,
        ),
        Spacer(1, 0.4 * cm),
        Paragraph(
            f"<b>{emp['prenom']} {emp['nom']}</b> "
            f"(matricule {emp['matricule']}, n° AVS {emp.get('no_avs', '—')})",
            sn,
        ),
        Spacer(1, 0.4 * cm),
        Paragraph(
            f"est employé(e) au sein de notre entreprise depuis le "
            f"<b>{datetime.strptime(emp['date_embauche'], '%Y-%m-%d').strftime('%d/%m/%Y')}</b>, "
            f"en qualité de <b>{emp['poste']}</b>, au sein du service {emp['service']} "
            f"(canton de {emp['canton']}).",
            sn,
        ),
    ]
    if motif:
        story += [
            Spacer(1, 0.4 * cm),
            Paragraph(f"<b>Motif :</b> {motif}", sn),
        ]
    story += [
        Spacer(1, 0.5 * cm),
        Paragraph(
            "La présente attestation est délivrée à l'intéressé(e) pour servir et "
            "valoir ce que de droit.",
            sn,
        ),
        Spacer(1, 2 * cm),
        Paragraph("_______________________________", sn),
        Paragraph("Léa MARTIN<br/>Directrice des Ressources Humaines", sn),
    ]
    doc.build(story)
    return f"✅ Attestation PDF générée : {pdf_path.name}{_marqueur(pdf_path)}"


# =====================================================================
# DOCX — attestation Word modifiable
# =====================================================================
def exporter_attestation_docx(matricule: str, motif: str = "") -> str:
    """Génère une attestation de travail au format Word modifiable."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return "❌ Lance : pip install python-docx"

    emp = next((e for e in _charger_employes() if e["matricule"].upper() == matricule.upper()), None)
    if not emp:
        return f"Matricule introuvable : {matricule}"

    docx_path = EXPORTS_DIR / f"attestation_{matricule}_{date.today().strftime('%Y%m%d')}.docx"
    doc = Document()
    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)

    # En-tête
    p = doc.add_paragraph()
    p.add_run(ENTREPRISE["nom"]).bold = True
    doc.add_paragraph(f"{ENTREPRISE['adresse']}\n{ENTREPRISE['npa']} {ENTREPRISE['ville']}")

    p = doc.add_paragraph()
    p.add_run(f"Fait à {ENTREPRISE['ville']}, le {date.today().strftime('%d/%m/%Y')}").italic = True

    doc.add_paragraph()

    # Titre
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titre.add_run("ATTESTATION DE TRAVAIL")
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph()
    doc.add_paragraph(
        f"Je soussignée, en qualité de Directrice des Ressources Humaines de la société "
        f"{ENTREPRISE['nom']}, atteste par la présente que :"
    )
    doc.add_paragraph(
        f"{emp['prenom']} {emp['nom']} (matricule {emp['matricule']}, n° AVS {emp.get('no_avs', '—')})"
    ).runs[0].bold = True

    doc.add_paragraph(
        f"est employé(e) au sein de notre entreprise depuis le "
        f"{datetime.strptime(emp['date_embauche'], '%Y-%m-%d').strftime('%d/%m/%Y')}, "
        f"en qualité de {emp['poste']}, au sein du service {emp['service']} "
        f"(canton de {emp['canton']})."
    )

    if motif:
        doc.add_paragraph(f"Motif : {motif}")

    doc.add_paragraph(
        "La présente attestation est délivrée à l'intéressé(e) pour servir et "
        "valoir ce que de droit."
    )

    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph("_______________________________")
    doc.add_paragraph("Léa MARTIN\nDirectrice des Ressources Humaines")

    doc.save(str(docx_path))
    return f"✅ Attestation Word générée : {docx_path.name}{_marqueur(docx_path)}"


# =====================================================================
# XLSX — exports Excel
# =====================================================================
def exporter_employes_xlsx() -> str:
    """Exporte la liste complète des employés en Excel."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        return "❌ Lance : pip install openpyxl"

    employes = _charger_employes()
    if not employes:
        return "Aucun employé à exporter."

    xlsx_path = EXPORTS_DIR / f"employes_{date.today().strftime('%Y%m%d')}.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Employés"

    # En-têtes
    headers = list(employes[0].keys())
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="4F46E5", end_color="4F46E5", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # Données
    for emp in employes:
        ws.append([emp[k] for k in headers])

    # Largeurs auto
    for col in ws.columns:
        max_l = max(len(str(c.value)) for c in col if c.value is not None)
        ws.column_dimensions[col[0].column_letter].width = min(max_l + 2, 35)

    ws.freeze_panes = "A2"
    wb.save(str(xlsx_path))
    return f"✅ Excel des employés généré : {xlsx_path.name} ({len(employes)} lignes){_marqueur(xlsx_path)}"


def exporter_factures_xlsx(statut: str = "tout") -> str:
    """Exporte les factures en Excel (avec totaux)."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        return "❌ Lance : pip install openpyxl"

    fact_csv = Path(__file__).parent.parent / "donnees" / "factures.csv"
    if not fact_csv.exists():
        return "Aucune facture."
    with open(fact_csv, encoding="utf-8") as f:
        factures = list(csv.DictReader(f))

    if statut != "tout":
        factures = [f for f in factures if f["statut"].lower() == statut.lower()]
    if not factures:
        return f"Aucune facture avec statut '{statut}'."

    xlsx_path = EXPORTS_DIR / f"factures_{statut}_{date.today().strftime('%Y%m%d')}.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = f"Factures ({statut})"

    headers = ["N°", "Client", "Émission", "Échéance", "HT (CHF)", "TVA %",
               "TVA (CHF)", "TTC (CHF)", "Description", "Statut"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="4F46E5", end_color="4F46E5", fill_type="solid")
        cell.alignment = Alignment(horizontal="center")

    total_ht = total_tva = total_ttc = 0.0
    for f in factures:
        client = _trouver_client(f["id_client"])
        nom_cli = client["nom"] if client else f["id_client"]
        ht, tva, ttc = float(f["montant_ht"]), float(f["montant_tva"]), float(f["montant_ttc"])
        total_ht += ht; total_tva += tva; total_ttc += ttc
        ws.append([
            f["numero"], nom_cli, f["date_emission"], f["date_echeance"],
            ht, float(f["taux_tva"]), tva, ttc, f["description"], f["statut"]
        ])

    # Ligne totaux
    n_row = ws.max_row + 1
    ws.cell(row=n_row, column=4, value="TOTAUX").font = Font(bold=True)
    ws.cell(row=n_row, column=5, value=total_ht).font = Font(bold=True)
    ws.cell(row=n_row, column=7, value=total_tva).font = Font(bold=True)
    ws.cell(row=n_row, column=8, value=total_ttc).font = Font(bold=True)
    for c in (5, 7, 8):
        ws.cell(row=n_row, column=c).fill = PatternFill(
            start_color="EEF2FF", end_color="EEF2FF", fill_type="solid"
        )

    # Format monétaire
    for col in (5, 7, 8):
        for row in range(2, ws.max_row + 1):
            ws.cell(row=row, column=col).number_format = '#,##0.00 "CHF"'

    # Largeurs
    for i, h in enumerate(headers, 1):
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = max(15, len(h) + 4)

    ws.freeze_panes = "A2"
    wb.save(str(xlsx_path))
    return (
        f"✅ Excel généré : {xlsx_path.name}\n"
        f"   {len(factures)} factures, total TTC {total_ttc:,.2f} CHF".replace(",", "'")
        + _marqueur(xlsx_path)
    )
